"""Gera um .docx com a Ordem de Culto do Encontro Real (20/06/2026).

Layout: duas cópias idênticas por página, separadas por uma linha de corte
tracejada, para impressão e recorte ao meio.

Execução (a partir da raiz do repositório):
    python3 referencia/scripts/gerar_ordem_culto.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

TITULO = "Ordem de Culto — Encontro Real"
SUBTITULO = "20/06/2026"

LINHAS = [
    ("1", "Abertura / Boas-vindas", "Recepção da igreja e equipes (13h)", "Lucas e Ana"),
    ("2", "Entrada das bandeiras", "Nacional, ER e MR — fundo instrumental", "Equipe de cerimonial"),
    ("3", "Compromissos MR", "Divisa, ideais, pacto e hino", "Equipe MR"),
    ("4", "Compromissos ER", "Divisa, pátria, bandeira cristã, dos ER e hino", "Equipe ER"),
    ("5", "Divisão para as provas", "Slide com nome da prova e responsáveis", "Ana Paula"),
    ("6", "Kahoot no templo", "Dinâmica interativa", "Thalyta"),
    (
        "7",
        "Louvor — Vila Garrido",
        "1) Vida com Deus\n2) Qualquer coisa que eu fizer\n3) Diante de Ti",
        "Juan (ministro)",
    ),
    (
        "8",
        "Palavra",
        "Pregação + entrega da lembrança ao final",
        "Pr. Jullyander • Julia (chamada / oração / lembrança)",
    ),
    (
        "9",
        "Louvor (pós-palavra)",
        "1) Toda sorte de bênçãos\n2) Que bonito é\n(ajustar conforme a premiação)",
        "Equipe de louvor",
    ),
    ("10", "Premiação", "Entrega de prêmios das provas", "Todos"),
    ("11", "Oração final e agradecimentos", "Encerramento do culto", "Todos"),
]

CABECALHOS = ["#", "Momento", "Detalhes", "Responsáveis"]
# Larguras das colunas (somam a área útil de uma página A4 retrato com margens reduzidas).
LARGURAS_CM = [1.0, 4.2, 7.5, 5.3]


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color: str = "888888", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tc_pr.append(borders)


def _adicionar_titulo(doc: Document) -> None:
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(TITULO)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
    titulo.paragraph_format.space_before = Pt(0)
    titulo.paragraph_format.space_after = Pt(0)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run(SUBTITULO)
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after = Pt(4)


def _adicionar_tabela(doc: Document) -> None:
    tabela = doc.add_table(rows=1 + len(LINHAS), cols=len(CABECALHOS))
    tabela.autofit = False
    tabela.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, largura in enumerate(LARGURAS_CM):
        for row in tabela.rows:
            row.cells[idx].width = Cm(largura)

    cabecalho = tabela.rows[0]
    cabecalho.height = Cm(0.7)
    cabecalho.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for col, texto in enumerate(CABECALHOS):
        cell = cabecalho.cells[col]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(cell, "1F2A44")
        _set_cell_borders(cell)
        paragrafo = cell.paragraphs[0]
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragrafo.add_run(texto)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for linha_idx, linha in enumerate(LINHAS, start=1):
        cor_fundo = "F4F6FB" if linha_idx % 2 == 0 else "FFFFFF"
        row = tabela.rows[linha_idx]
        row.height = Cm(0.55)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for col, conteudo in enumerate(linha):
            cell = row.cells[col]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_shading(cell, cor_fundo)
            _set_cell_borders(cell, color="C8CCD6")
            cell.text = ""
            partes = str(conteudo).split("\n")
            for parte_idx, parte in enumerate(partes):
                paragrafo = cell.paragraphs[0] if parte_idx == 0 else cell.add_paragraph()
                paragrafo.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if col == 0 else WD_ALIGN_PARAGRAPH.LEFT
                )
                paragrafo.paragraph_format.space_before = Pt(0)
                paragrafo.paragraph_format.space_after = Pt(0)
                run = paragrafo.add_run(parte)
                run.font.size = Pt(9.5)
                if col == 1:
                    run.bold = True
                if col == 0:
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _adicionar_separador_de_corte(doc: Document) -> None:
    paragrafo = doc.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo.paragraph_format.space_before = Pt(10)
    paragrafo.paragraph_format.space_after = Pt(10)
    run = paragrafo.add_run("✂ — — — — — — — — — — — — — — — — — — — — — — — — — — — — — — —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _adicionar_quebra_de_pagina(doc: Document) -> None:
    paragrafo = doc.add_paragraph()
    paragrafo.paragraph_format.space_before = Pt(0)
    paragrafo.paragraph_format.space_after = Pt(0)
    run = paragrafo.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _adicionar_bloco(doc: Document) -> None:
    _adicionar_titulo(doc)
    _adicionar_tabela(doc)


def gerar(caminho_saida: Path, paginas: int = 1) -> Path:
    doc = Document()

    for secao in doc.sections:
        secao.top_margin = Cm(1.2)
        secao.bottom_margin = Cm(1.2)
        secao.left_margin = Cm(1.5)
        secao.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for pagina_idx in range(paginas):
        _adicionar_bloco(doc)
        _adicionar_separador_de_corte(doc)
        _adicionar_bloco(doc)
        if pagina_idx < paginas - 1:
            _adicionar_quebra_de_pagina(doc)

    caminho_saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(caminho_saida)
    return caminho_saida


if __name__ == "__main__":
    raiz = Path(__file__).resolve().parent.parent
    destino = raiz / "Ordem_Culto_Encontro_Real_2026-06-20.docx"
    final = gerar(destino, paginas=1)
    print(f"Gerado: {final}")
